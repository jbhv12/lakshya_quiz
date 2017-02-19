import docx,random
from quiz_data2 import d,l
random.seed(100)
doc = docx.Document()
for p in range(5):
    random.shuffle(l)
    f = open('slot2 '+str(p+1),'w')
    doc.add_paragraph('question set id: '+str(p+1))
    para = doc.add_paragraph('')
    for i,n in zip(l,range(len(l))):
        que = list(d.keys())[i] #txt word
        ans = d[que]
        optl = [ans[0]] + ans[1]
        random.shuffle(optl)
        opt_str=''
        for o,letter in zip(optl,list('abcd')):
            opt_str += letter+'. '+o+' '
        print(que)
        f.write(str(n+1) + '. ' + que + ans[0] + '\n')
        para.add_run(str(n+1) + '. ' + que + '\n')
        para.add_run(opt_str+'\n')
    para.runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    f.close()
    
doc.save('slot2_5variations.docx')

